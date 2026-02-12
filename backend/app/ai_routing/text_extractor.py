from PyPDF2 import PdfReader
import docx
import re
import logging

logger = logging.getLogger(__name__)


# ===============================
# PDF TEXT EXTRACTION
# ===============================
def extract_text_from_pdf(path: str) -> str:
    text_chunks: list[str] = []

    try:
        reader = PdfReader(path)

        for page in reader.pages:
            try:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_chunks.append(page_text)
            except Exception:
                continue

    except Exception as e:
        logger.exception(f"PDF extract failed: {path}")

    return normalize_text("\n".join(text_chunks))


# ===============================
# DOCX TEXT EXTRACTION (FULL)
# ===============================
def extract_text_from_docx(path: str) -> str:
    text_chunks: list[str] = []
    seen: set[str] = set()

    try:
        doc = docx.Document(path)

        # paragraphs
        for p in doc.paragraphs:
            text = p.text.strip()
            if text and text not in seen:
                seen.add(text)
                text_chunks.append(text)

        # tables (deduplicated)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in seen:
                        seen.add(cell_text)
                        text_chunks.append(cell_text)

    except Exception:
        logger.exception(f"DOCX extract failed: {path}")

    return normalize_text("\n".join(text_chunks))


# ===============================
# NORMALIZE TEXT (CRITICAL)
# ===============================
def normalize_text(text: str) -> str:
    if not text:
        return ""

    # collapse whitespace
    text = re.sub(r"\s+", " ", text)

    # normalize date separators
    text = re.sub(r"\s*/\s*", "/", text)
    text = re.sub(r"\s*-\s*", "-", text)

    return text.strip()


# ===============================
# MAIN ENTRY
# ===============================
def extract_text(file_path: str, file_type: str) -> str:
    if not file_path or not file_type:
        return ""

    file_type = file_type.lower()

    if "pdf" in file_type:
        return extract_text_from_pdf(file_path)

    if "word" in file_type or "docx" in file_type:
        return extract_text_from_docx(file_path)

    logger.warning(f"Unsupported file type: {file_type}")
    return ""
