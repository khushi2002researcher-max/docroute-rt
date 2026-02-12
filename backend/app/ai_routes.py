# app/ai_routes.py

from fastapi import APIRouter, UploadFile, File, Depends, HTTPException, Form, BackgroundTasks
from sqlalchemy.orm import Session
from fastapi.responses import JSONResponse, FileResponse
from pathlib import Path
from tempfile import NamedTemporaryFile
from io import BytesIO
import os

from app.database import get_db
from app.auth.utils import get_current_user
from app.auth.models import User, AIDocument
from app.summarizer import generate_summary
from app.tagger import generate_tags

from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError
from docx import Document

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


ai_router = APIRouter(prefix="/ai", tags=["AI"])
ALLOWED_EXTENSIONS = {"txt", "pdf", "docx"}

BASE_DIR = Path(__file__).resolve().parent
FONT_PATH = BASE_DIR / "fonts" / "NotoSansDevanagari-Regular.ttf"

# ✅ Safe Font Registration
if FONT_PATH.exists():
    if "NotoDeva" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("NotoDeva", str(FONT_PATH)))
else:
    print("⚠ Hindi font not found. Hindi PDF may not render correctly.")


def safe_text(text: str) -> str:
    if not text or not text.strip():
        return "[No readable text detected]"
    return text[:2000]


# ---------------------------
# Analyze file
# ---------------------------
@ai_router.post("/analyze-file")
async def analyze_file(
    upload_file: UploadFile = File(...),
    language: str = Form(...),
    length: str = Form("short"),
    format: str = Form("paragraph"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    contents = await upload_file.read()

    # ✅ File size limit (5MB)
    if len(contents) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 5MB)")

    # ✅ Normalize language
    language = language.lower()
    if language in ["en", "english"]:
        language = "english"
    elif language in ["hi", "hindi"]:
        language = "hindi"
    else:
        raise HTTPException(status_code=400, detail="Only English and Hindi supported")

    ext = upload_file.filename.split(".")[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    text = ""

    if ext == "txt":
        text = contents.decode(errors="ignore")

    elif ext == "pdf":
        try:
            reader = PdfReader(BytesIO(contents))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except PdfReadError:
            raise HTTPException(status_code=400, detail="Invalid or corrupted PDF")

    elif ext == "docx":
        doc = Document(BytesIO(contents))
        text = "\n".join(p.text for p in doc.paragraphs)

    text = safe_text(text)

    summary, processed_text = generate_summary(
        text=text,
        length=length,
        format=format,
        language=language,
    )

    tags = generate_tags(
        text=processed_text,
        language=language,
    )

    record = AIDocument(
        user_id=user.id,
        file_name=upload_file.filename,
        input_text=text,
        language=language,
        length=length,
        format=format,
        summary=summary,
        tags=tags,
    )

    db.add(record)
    db.commit()
    db.refresh(record)

    return {"summary": summary, "tags": tags}


# ---------------------------
# History
# ---------------------------
@ai_router.get("/history")
async def get_history(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    docs = db.query(AIDocument).filter(AIDocument.user_id == user.id).all()
    return [
        {
            "id": d.id,
            "file_name": d.file_name,
            "summary": d.summary,
            "tags": d.tags,
            "language": d.language,
            "length": d.length,
            "format": d.format,
            "created_at": d.created_at,
        }
        for d in docs
    ]


# ---------------------------
# Delete
# ---------------------------
@ai_router.delete("/{doc_id}")
async def delete_document(doc_id: int, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    doc = db.query(AIDocument).filter(
        AIDocument.id == doc_id,
        AIDocument.user_id == user.id
    ).first()

    if not doc:
        raise HTTPException(404, "Document not found")

    db.delete(doc)
    db.commit()

    return {"message": "Deleted successfully"}


# ---------------------------
# Download TXT
# ---------------------------
@ai_router.get("/download/{doc_id}")
async def download_txt(
    doc_id: int,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    doc = db.query(AIDocument).filter(
        AIDocument.id == doc_id,
        AIDocument.user_id == user.id
    ).first()

    if not doc:
        raise HTTPException(404, "Document not found")

    temp_file = NamedTemporaryFile(delete=False, suffix=".txt")
    temp_file.write(doc.input_text.encode())
    temp_file.close()

    background_tasks.add_task(os.remove, temp_file.name)

    return FileResponse(temp_file.name, filename=doc.file_name, media_type="text/plain")


# ---------------------------
# Download PDF
# ---------------------------
@ai_router.get("/download/pdf/{doc_id}")
async def download_pdf(
    doc_id: int,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    doc = db.query(AIDocument).filter(
        AIDocument.id == doc_id,
        AIDocument.user_id == user.id
    ).first()

    if not doc:
        raise HTTPException(404, "Document not found")

    temp_file = NamedTemporaryFile(delete=False, suffix=".pdf")

    styles = getSampleStyleSheet()
    style = styles["Normal"]

    if doc.language == "hindi":
        style.fontName = "NotoDeva"

    pdf = SimpleDocTemplate(temp_file.name, pagesize=A4)
    content = [
        Paragraph("<b>AI Summary</b>", style),
        Paragraph(doc.summary, style),
        Paragraph("<br/><b>Tags:</b>", style),
        Paragraph(", ".join(doc.tags or []), style),
    ]

    pdf.build(content)

    background_tasks.add_task(os.remove, temp_file.name)

    return FileResponse(
        temp_file.name,
        filename=doc.file_name.replace(".", "_") + ".pdf",
        media_type="application/pdf"
    )


# ---------------------------
# Download DOCX
# ---------------------------
@ai_router.get("/download/docx/{doc_id}")
async def download_docx(
    doc_id: int,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    doc = db.query(AIDocument).filter(
        AIDocument.id == doc_id,
        AIDocument.user_id == user.id
    ).first()

    if not doc:
        raise HTTPException(404, "Document not found")

    temp_file = NamedTemporaryFile(delete=False, suffix=".docx")
    document = Document()

    document.add_heading("AI Summary", level=1)
    document.add_paragraph(doc.summary)

    document.add_heading("Tags", level=2)
    document.add_paragraph(", ".join(doc.tags or []))

    document.save(temp_file.name)

    background_tasks.add_task(os.remove, temp_file.name)

    return FileResponse(
        temp_file.name,
        filename=doc.file_name.replace(".", "_") + ".docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
