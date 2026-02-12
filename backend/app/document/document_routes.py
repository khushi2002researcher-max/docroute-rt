from fastapi import (
    APIRouter,
    Depends,
    HTTPException,
    UploadFile,
    File,
    Form,
    Request,
    
)
from sqlalchemy import func
from sqlalchemy.orm import Session
from pathlib import Path

from typing import Optional
from uuid import uuid4
import difflib


from app.database import get_db
from app.auth.models import User
from app.document.D_models import (
    Document,
    DocumentVersion,
    AuditLog,
    DocumentShareLink,
    DocumentDownloadLog,
)
from fastapi.responses import FileResponse
import tempfile
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from datetime import datetime, timezone

from app.auth.utils import (
    get_current_user,
    generate_tracking_id,
    build_client_info,
    
    is_share_link_valid,
    verify_share_link_password,
    check_share_permission,
    hash_password,
    revoke_share_link,   # ✅ ADD THIS
)



router = APIRouter(prefix="/documents", tags=["Documents"])

# =====================================================
# UPLOAD DIRECTORY
# =====================================================
UPLOAD_DIR = Path("uploads/documents")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# =====================================================
# LIVE VIEWERS (WEBSOCKET STORAGE)
# =====================================================

# =====================================================
# ACCESS CHECK
# =====================================================
def check_access(user: User, document: Document, required: str):
    if document.owner_id == user.id:
        return

    access = next(
        (
            a for a in document.access_list
            if a.user_id == user.id and a.revoked_at is None
        ),
        None,
    )

    if not access:
        raise HTTPException(status_code=403, detail="No access to document")

    if required == "edit" and access.permission != "edit":
        raise HTTPException(status_code=403, detail="Edit permission required")

# =====================================================
# CREATE DOCUMENT (NEW + UPLOAD)
# =====================================================
@router.post("/create")
def create_document(
    file_name: str = Form(...),
    output_type: str = Form(...),
    content: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
    
):
    if not file_name.strip():
        raise HTTPException(status_code=422, detail="File name required")

    
    if file is None and (content is None or content.strip() == ""):
        raise HTTPException(
            status_code=422,
            detail="Either file or content must be provided",
        )

    if output_type not in ("docx", "pdf"):
        raise HTTPException(
            status_code=422,
            detail="Invalid output_type (must be docx or pdf)",
        )

    final_content = ""

    if file:
        filename = file.filename.lower()

        if filename.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(file.file)
            final_content = "\n".join(p.text for p in doc.paragraphs)

        elif filename.endswith(".pdf"):
            import PyPDF2
            reader = PyPDF2.PdfReader(file.file)
            final_content = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type (only PDF or DOCX)",
            )
    else:
        final_content = content.strip()

    stored_file = f"{uuid4().hex}.{output_type}"

    document = Document(
        file_name=file_name.strip(),
        content=final_content,
        owner_id=user.id,
        tracking_id=generate_tracking_id(),
        file_type=output_type,
        stored_file_name=stored_file,   # ✅ REQUIRED FIX
    )
    


    db.add(document)
    db.commit()
    db.refresh(document)

    return {
        "id": document.id,
        "file_name": document.file_name,
        "content": document.content,
        "tracking_id": document.tracking_id,
        "last_updated_at": document.last_updated_at,
    }

# =====================================================
# SHARE DOCUMENT
# =====================================================
@router.post("/{doc_id}/share")
def create_share_link(
    doc_id: int,
    permission: str = Form("view"),
    password: Optional[str] = Form(None),
    expires_in_minutes: int = Form(1440),  # ignored now
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = db.query(Document).filter(Document.id == doc_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    if permission not in ("view", "edit", "download"):
        raise HTTPException(status_code=400, detail="Invalid permission")


    link = DocumentShareLink(
        document_id=document.id,
        token=uuid4().hex,
        permission=permission,
        password_hash=hash_password(password) if password else None,
        created_by=current_user.id,
        expires_at=None,          # ✅ HERE — NO EXPIRY
        is_active=True,
        opened_count=0,
        revoked_at=None,
    )
    

    db.add(link)
    db.commit()

    return {
        "token": link.token,
        "expires_at": link.expires_at,  # will be null
    }


# =====================================================
# OPEN SHARED DOCUMENT
# =====================================================
@router.get("/share/{token}")
def open_shared_document(
    token: str,
    request: Request,
    password: Optional[str] = None,
    db: Session = Depends(get_db),
):
    link = db.query(DocumentShareLink).filter_by(token=token).first()
    if not link or not is_share_link_valid(link):
        raise HTTPException(status_code=410, detail="Link expired")

    verify_share_link_password(link, password)
    check_share_permission(link, "view")

    document = db.query(Document).filter(Document.id == link.document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # 🔥 FIX: track opens
    now = datetime.now(timezone.utc)

    link.opened_count += 1  
    link.last_opened_at = now

    watermark = {
        "text": f"Shared | {request.client.host} | {now.isoformat()}",
        "opacity": 0.15,
    }


    db.add(
        AuditLog(
            module="documents",
            action="open_shared",
            event_type="view",
            document_id=document.id,
            client_info=build_client_info(request),
        )
    )

    db.commit()

    return {
    "document": {
        "id": document.id,
        "file_name": document.file_name,
        "content": document.content,   # ✅ THIS FIXES PREVIEW
        "tracking_id": document.tracking_id,
        "last_updated_at": document.last_updated_at,
    },
    "watermark": watermark,
    "permission": link.permission,
}


# =====================================================
# LIST DOCUMENTS
# =====================================================
@router.get("/")
def list_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    documents = (
        db.query(Document)
        .filter(
            Document.is_deleted == False,
            Document.owner_id == current_user.id,
        )
        .order_by(Document.last_updated_at.desc())
        .all()
    )

    result = []

    for d in documents:
        shared_opens = (
            db.query(DocumentShareLink)
            .filter(DocumentShareLink.document_id == d.id)
            .with_entities(func.coalesce(func.sum(DocumentShareLink.opened_count), 0))
            .scalar()
        )

        result.append({
            "id": d.id,
            "file_name": d.file_name,
            "tracking_id": d.tracking_id,
            "file_type": d.file_type,
            "last_updated_at": d.last_updated_at,
            "total_opens": shared_opens or 0,
        })

    return result


# =====================================================
# PREVIEW DOCUMENT
# =====================================================
@router.get("/{doc_id}")
def preview_document(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = db.query(Document).filter(
        Document.id == doc_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    check_access(current_user, document, "view")

    document.last_viewed_at = datetime.now(timezone.utc)

    document.last_viewed_by = current_user.id

    db.add(
        AuditLog(
            module="documents",
            action="view",
            event_type="view",
            document_id=document.id,
            performed_by=current_user.id,
            client_info=build_client_info(request),
        )
    )
    owner_views = (
        db.query(AuditLog)
        .filter(
            AuditLog.document_id == document.id,
            AuditLog.action == "view",
        )
        .count()
    )

# Shared link opens
    shared_opens = (
        db.query(DocumentShareLink)
        .filter(DocumentShareLink.document_id == document.id)
        .with_entities(func.coalesce(func.sum(DocumentShareLink.opened_count), 0))
        .scalar()
    )

    total_opens = owner_views + (shared_opens or 0)

    db.commit()

    return {
        "id": document.id,
        "file_name": document.file_name,
        "content": document.content,
        "tracking_id": document.tracking_id,
        "last_updated_at": document.last_updated_at,
        "total_opens": total_opens,   # ✅ ADD THIS
    }

# =====================================================
# UPDATE DOCUMENT (🔥 FIXED)
# =====================================================
@router.put("/{doc_id}/update")
def update_document(
    doc_id: int,
    content: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = db.query(Document).filter(
        Document.id == doc_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    check_access(current_user, document, "edit")

    if content is None or content.strip() == "":
        raise HTTPException(
            status_code=422,
            detail="Content cannot be empty",
        )

    previous_content = document.content or ""
    document.content = content.strip()

    version_number = (
        db.query(DocumentVersion)
        .filter(DocumentVersion.document_id == document.id)
        .count()
        + 1
    )

    diff_text = "\n".join(
        difflib.ndiff(
            previous_content.splitlines(),
            document.content.splitlines(),
        )
    )

    db.add(
        DocumentVersion(
            document_id=document.id,
            version_number=version_number,
            content=document.content,
            diff=diff_text,
            created_by=current_user.id,
        )
    )

    db.add(
        AuditLog(
            module="documents",
            action="update",
            event_type="edit",
            document_id=document.id,
            performed_by=current_user.id,
        )
    )

    db.commit()

    return {
        "id": document.id,
        "file_name": document.file_name,
        "content": document.content,
        "tracking_id": document.tracking_id,
        "last_updated_at": document.last_updated_at,
    }



# =====================================================
# DOWNLOAD TRACKING (🔥 FIXED)
# =====================================================
@router.get("/{doc_id}/download")
def track_download(
    doc_id: int,
    format: str,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if format not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="Invalid download format")

    document = db.query(Document).filter(
        Document.id == doc_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    check_access(current_user, document, "view")

    db.add(
        DocumentDownloadLog(
            document_id=document.id,
            downloaded_by=current_user.id,
            format=format,
            client_info=build_client_info(request),
        )
    )

    db.commit()
    return {"status": "logged"}

# =====================================================
# WEBSOCKET — LIVE VIEWERS
# =====================================================


# =====================================================
# DELETE DOCUMENT (SOFT DELETE)
# =====================================================
@router.delete("/{doc_id}")
def delete_document(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = db.query(Document).filter(
        Document.id == doc_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    if document.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Only owner can delete document")

    document.is_deleted = True

    db.add(
        AuditLog(
            module="documents",
            action="delete",
            event_type="delete",
            document_id=document.id,
            performed_by=current_user.id,
            client_info=build_client_info(request),
        )
    )

    db.commit()

    return {"status": "deleted"}


# =====================================================
# EDIT SHARED DOCUMENT (NEW)
# =====================================================
# =====================================================
# EDIT DOCUMENT VIA SHARE LINK
# =====================================================
@router.put("/share/{token}/edit")
def edit_shared_document(
    token: str,
    content: str = Form(...),
    request: Request = None,
    db: Session = Depends(get_db),
):
    link = db.query(DocumentShareLink).filter_by(token=token).first()
    if not link or not is_share_link_valid(link):
        raise HTTPException(status_code=410, detail="Link expired")

    check_share_permission(link, "edit")

    document = db.query(Document).filter(Document.id == link.document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    if not content.strip():
        raise HTTPException(status_code=422, detail="Content cannot be empty")

    # 🔹 Save previous content
    previous_content = document.content or ""
    document.content = content.strip()

    # 🔹 Create version
    version_number = (
        db.query(DocumentVersion)
        .filter(DocumentVersion.document_id == document.id)
        .count()
        + 1
    )

    diff_text = "\n".join(
        difflib.ndiff(
            previous_content.splitlines(),
            document.content.splitlines(),
        )
    )

    db.add(
        DocumentVersion(
            document_id=document.id,
            version_number=version_number,
            content=document.content,
            diff=diff_text,
            created_by=link.created_by,  # shared editor
        )
    )

    # 🔹 Audit log
    db.add(
        AuditLog(
            module="documents",
            action="shared_edit",
            event_type="edit",
            document_id=document.id,
            client_info=build_client_info(request),
        )
    )

    db.commit()

    return {"status": "updated"}



# =====================================================
# DOWNLOAD DOCUMENT VIA SHARE LINK (DOCX OR PDF)
# =====================================================
@router.get("/share/{token}/download")
def download_shared_document(
    token: str,
    format: str = "docx",  # docx | pdf
    request: Request = None,
    db: Session = Depends(get_db),
):
    link = db.query(DocumentShareLink).filter_by(token=token).first()
    if not link or not is_share_link_valid(link):
        raise HTTPException(status_code=410, detail="Link expired")

    check_share_permission(link, "download")

    if format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="Invalid format")

    document = db.query(Document).filter(Document.id == link.document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # 🔹 Audit log
    db.add(
        AuditLog(
            module="documents",
            action="shared_download",
            event_type="download",
            document_id=document.id,
            client_info=build_client_info(request),
        )
    )
    db.commit()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{format}")



    # ---------- DOCX ----------
    if format == "docx":
        doc = DocxDocument()
        for line in (document.content or "").splitlines():
            doc.add_paragraph(line)

        tmp.close()
        doc.save(tmp.name)

        return FileResponse(
            tmp.name,
            filename=f"{document.file_name}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # ---------- PDF ----------
    c = canvas.Canvas(tmp.name, pagesize=A4)
    text = c.beginText(40, 800)

    for line in (document.content or "").splitlines():
        text.textLine(line)

    c.drawText(text)
    c.showPage()
    c.save()
    tmp.close()

    return FileResponse(
        tmp.name,
        filename=f"{document.file_name}.pdf",
        media_type="application/pdf",
    )


@router.post("/share/{token}/revoke")
def revoke_share(
    token: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    link = db.query(DocumentShareLink).filter_by(token=token).first()
    if not link:
        raise HTTPException(status_code=404, detail="Link not found")

    if link.created_by != current_user.id:
        raise HTTPException(status_code=403, detail="Not allowed")

    revoke_share_link(link)  # ✅ uses fixed helper

    db.commit()
    return {"status": "revoked"}


@router.get("/{doc_id}/versions")
def list_versions(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = db.query(Document).filter(
        Document.id == doc_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    check_access(current_user, document, "view")

    versions = (
        db.query(DocumentVersion)
        .filter(DocumentVersion.document_id == doc_id)
        .order_by(DocumentVersion.version_number.desc())
        .all()
    )

    return [
        {
            "id": v.id,
            "version_number": v.version_number,
            "created_at": v.created_at,
            "created_by": v.created_by,
        }
        for v in versions
    ]

@router.get("/{doc_id}/versions/{version_id}/diff")
def version_diff(
    doc_id: int,
    version_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = db.query(Document).filter(
        Document.id == doc_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    check_access(current_user, document, "view")

    version = (
        db.query(DocumentVersion)
        .filter(
            DocumentVersion.id == version_id,
            DocumentVersion.document_id == doc_id,
        )
        .first()
    )

    if not version:
        raise HTTPException(status_code=404, detail="Version not found")

    # Find previous version
    prev_version = (
        db.query(DocumentVersion)
        .filter(
            DocumentVersion.document_id == doc_id,
            DocumentVersion.version_number == version.version_number - 1,
        )
        .first()
    )

    old_text = prev_version.content if prev_version else ""
    new_text = version.content or ""

    diff = list(difflib.ndiff(
        old_text.splitlines(),
        new_text.splitlines(),
    ))

    # 🔹 Summary
    added = [l for l in diff if l.startswith("+ ")]
    removed = [l for l in diff if l.startswith("- ")]

    summary = {
        "lines_added": len(added),
        "lines_removed": len(removed),
        "total_changes": len(added) + len(removed),
    }

    return {
        "previous_version": prev_version.version_number if prev_version else None,
        "current_version": version.version_number,
        "diff": diff,
        "summary": summary,
    }


@router.get("/{doc_id}/audit-logs")
def get_document_audit_logs(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = db.query(Document).filter(
        Document.id == doc_id,
        Document.is_deleted == False,
    ).first()

    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # 🔐 Only owner can view full audit
    if document.owner_id != current_user.id:
        raise HTTPException(
            status_code=403,
            detail="Only owner can view audit logs",
        )

    logs = (
        db.query(AuditLog)
        .filter(AuditLog.document_id == doc_id)
        .order_by(AuditLog.created_at.desc())
        .all()
    )

    return [
        {
            "id": log.id,
            "action": log.action,
            "event_type": log.event_type,
            "performed_by": (
                log.performed_by_user.email
                if log.performed_by_user
                else "Shared / Anonymous"
            ),
            "client_info": log.client_info,
            "created_at": log.created_at,
        }
        for log in logs
    ]


@router.delete("/audit-logs/{log_id}")
def delete_single_audit_log(
    log_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    log = db.query(AuditLog).filter(AuditLog.id == log_id).first()

    if not log:
        raise HTTPException(status_code=404, detail="Audit log not found")

    document = db.query(Document).filter(Document.id == log.document_id).first()

    # 🔐 Only document owner can delete logs
    if not document or document.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="Not allowed")

    db.delete(log)
    db.commit()

    return {"status": "audit log deleted"}



